import fitz
import pandas as pd
from docx import Document
import os

def extract_from_pdf(file_path: str) -> dict:
    texts = ""
    doc = fitz.open(file_path)
    for page_num in range(len(doc)):
        page = doc.load_page(page_num)
        texts += page.get_text() + "\n"

    metaData = {
        "fileName" : os.path.basename(file_path),
        "fileType" : 'Pdf',
        "numberOfPages" : len(doc)
    }
    return {"meta_data": metaData, "texts" : texts}

def extract_from_excel(file_path: str) -> dict:

    df = pd.read_excel(file_path)
    contents = df.to_string(index=False)
    metaData = {
        "fileName" : os.path.basename(file_path),
        "fileType" : 'Excel',
        "numberOfPages" : len(df)
    }
    return {"meta_data": metaData, "texts" : contents}

def extract_from_csv(file_path: str) -> dict:

    df = pd.read_csv(file_path)
    contents = df.to_string(index=False) # Don't add index column
    
    metaData = {
        "fileName" : os.path.basename(file_path),
        "fileType" : 'Csv',
        "numberOfPages" : len(df)
    }

    return {"meta_data": metaData, "texts" : contents}

def extract_from_word(file_path: str) -> dict:
    doc = Document(file_path)  # Initialize -> It will return as object
    
    # Extract all the texts from each paragraphs
    contents = '\n'.join([para.text for para in doc.paragraphs])
    
    metaData = {
        "fileName" : os.path.basename(file_path),
        "fileType" : 'Word',
        "numberOfPages" : len(doc.paragraphs)
    }
        
    return {"meta_data": metaData, "texts" : contents}

# Main Program to Process all files
def extract_text(file_path):
    ext = os.path.splitext(file_path)[1].lower()

    if ext == '.pdf':
        return extract_from_pdf(file_path=file_path)
    elif ext == '.csv':
        return extract_from_csv(file_path=file_path)
    elif ext == '.xlsx':
        return extract_from_excel(file_path=file_path)
    elif ext == '.docx':
        return extract_from_word(file_path=file_path)
    else:
        raise ValueError("Unsupported File.")